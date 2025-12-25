"""
Module d'extraction et de segmentation intelligente des documents d'entreprise
"""
import os
import re
import logging
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass, field
from enum import Enum

import pdfplumber
import PyPDF2
from docx import Document as DocxDocument

# Créer le logger pour ce module
logger = logging.getLogger(__name__)

# Essayer d'importer spaCy (optionnel)
try:
    import spacy
    SPACY_AVAILABLE = True
except ImportError:
    SPACY_AVAILABLE = False
    logger.warning("spaCy non disponible. Certaines fonctionnalités seront limitées.")


class DocumentType(Enum):
    """Types de documents supportés"""
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    UNKNOWN = "unknown"


@dataclass
class DocumentSection:
    """Représente une section structurée du document"""
    title: str
    content: str
    section_number: Optional[str] = None
    page_number: Optional[int] = None
    keywords: List[str] = field(default_factory=list)
    section_type: Optional[str] = None  # "procedure", "formulaire", "instruction"


class DocumentExtractor:
    """Extracteur intelligent de documents"""
    
    def __init__(self):
        """Initialise l'extracteur"""
        # Charger le modèle spaCy pour le français (si disponible)
        self.nlp = None
        if SPACY_AVAILABLE:
            try:
                self.nlp = spacy.load("fr_core_news_md")
            except OSError:
                logger.warning("Modèle spaCy non trouvé. Installez-le avec: python -m spacy download fr_core_news_md")
        
        # Patterns pour identifier les sections
        self.section_patterns = [
            r'^#+\s+(.+)$',  # Markdown headers
            r'^(\d+\.?\d*\.?\d*)\s+(.+)$',  # Numérotation (1., 1.1, 1.1.1)
            r'^([A-Z][A-Z\s]+)$',  # TITRES EN MAJUSCULES
            r'^\*\*(.+)\*\*$',  # **Titre en gras**
            r'^([A-Z]{2}-[A-Z]{2}-\d+)\s*-\s*(.+)$',  # FO-PR-01 - Titre
        ]
        
        # Mots-clés ISO pour classification
        self.iso_keywords = {
            "contexte": ["contexte", "organisation", "parties intéressées", "domaine d'application"],
            "leadership": ["leadership", "direction", "politique", "engagement", "responsabilité"],
            "planification": ["planification", "risques", "opportunités", "objectifs", "plan", "ordre de fabrication"],
            "support": ["ressources", "compétence", "formation", "communication", "information documentée", "composants", "stock"],
            "realisation": ["opérations", "exigences", "conception", "production", "prestation", "fabrication", "assemblage"],
            "evaluation": ["surveillance", "mesure", "analyse", "audit", "revue", "évaluation", "contrôle", "test", "conformité"],
            "amelioration": ["amélioration", "non-conformité", "action corrective", "amélioration continue"]
        }
    
    def detect_document_type(self, file_path: str) -> DocumentType:
        """Détecte le type de document"""
        extension = Path(file_path).suffix.lower()
        
        type_mapping = {
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.DOCX,
            '.doc': DocumentType.DOC
        }
        
        return type_mapping.get(extension, DocumentType.UNKNOWN)
    
    def extract_text(self, file_path: str) -> Tuple[str, Dict]:
        """
        Extrait le texte brut d'un document
        
        Returns:
            Tuple[str, Dict]: (texte_complet, métadonnées)
        """
        doc_type = self.detect_document_type(file_path)
        
        if doc_type == DocumentType.PDF:
            return self._extract_from_pdf(file_path)
        elif doc_type == DocumentType.DOCX:
            return self._extract_from_docx(file_path)
        else:
            raise ValueError(f"Type de document non supporté: {doc_type}")
    
    def _extract_from_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """Extrait le texte d'un PDF avec métadonnées"""
        full_text = []
        metadata = {
            "pages": 0,
            "file_name": Path(file_path).name,
            "file_type": "pdf"
        }
        
        try:
            with pdfplumber.open(file_path) as pdf:
                metadata["pages"] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        full_text.append(f"\n--- Page {page_num} ---\n{text}")
                        
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction PDF avec pdfplumber: {e}")
            # Fallback avec PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    metadata["pages"] = len(reader.pages)
                    
                    for page_num, page in enumerate(reader.pages, 1):
                        text = page.extract_text()
                        if text:
                            full_text.append(f"\n--- Page {page_num} ---\n{text}")
            except Exception as e2:
                logger.error(f"Erreur lors de l'extraction PDF avec PyPDF2: {e2}")
                raise
        
        return "\n".join(full_text), metadata
    
    def _extract_from_docx(self, file_path: str) -> Tuple[str, Dict]:
        """Extrait le texte d'un document Word avec structure"""
        doc = DocxDocument(file_path)
        
        full_text = []
        metadata = {
            "paragraphs": len(doc.paragraphs),
            "file_name": Path(file_path).name,
            "file_type": "docx"
        }
        
        for para in doc.paragraphs:
            if para.text.strip():
                # Préserver la structure (titres, etc.)
                if para.style.name.startswith('Heading'):
                    full_text.append(f"\n## {para.text}\n")
                else:
                    full_text.append(para.text)
        
        # Extraire aussi les tableaux
        for table in doc.tables:
            table_text = "\n[TABLEAU]\n"
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                table_text += row_text + "\n"
            full_text.append(table_text + "[FIN TABLEAU]\n")
        
        return "\n".join(full_text), metadata
    
    def segment_document(self, text: str, metadata: Dict) -> List[DocumentSection]:
        """
        Segmente intelligemment le document en sections logiques
        
        Args:
            text: Texte complet du document
            metadata: Métadonnées du document
            
        Returns:
            Liste de sections structurées
        """
        sections = []
        
        # Détecter les sections par patterns
        lines = text.split('\n')
        current_section = None
        current_content = []
        current_page = 1
        
        for line in lines:
            line_stripped = line.strip()
            
            # Détecter changement de page
            if '--- Page' in line:
                page_match = re.search(r'Page (\d+)', line)
                if page_match:
                    current_page = int(page_match.group(1))
                continue
            
            # Ignorer les lignes vides
            if not line_stripped:
                continue
            
            # Vérifier si c'est un titre de section
            is_section_title = False
            section_number = None
            title = None
            
            for pattern in self.section_patterns:
                match = re.match(pattern, line_stripped)
                if match:
                    # Sauvegarder la section précédente
                    if current_section and current_content:
                        current_section.content = '\n'.join(current_content)
                        sections.append(current_section)
                    
                    # Créer nouvelle section
                    groups = match.groups()
                    if len(groups) >= 2:
                        # Pattern avec numéro et titre (ex: "1.2 Titre" ou "FO-PR-01 - Titre")
                        section_number = groups[0]
                        title = groups[1]
                    else:
                        # Pattern avec titre seulement
                        title = groups[0]
                    
                    current_section = DocumentSection(
                        title=title,
                        content="",
                        section_number=section_number,
                        page_number=current_page
                    )
                    current_content = []
                    is_section_title = True
                    break
            
            # Ajouter la ligne au contenu
            if not is_section_title:
                current_content.append(line_stripped)
        
        # Ajouter la dernière section
        if current_section and current_content:
            current_section.content = '\n'.join(current_content)
            sections.append(current_section)
        
        # Si pas de sections détectées, créer une section globale
        if not sections:
            sections.append(DocumentSection(
                title=metadata.get("file_name", "Document complet"),
                content=text,
                page_number=1
            ))
        
        # Enrichir avec mots-clés et classification
        sections = self._classify_sections(sections)
        
        return sections
    
    def _classify_sections(self, sections: List[DocumentSection]) -> List[DocumentSection]:
        """Classifie les sections et extrait les mots-clés ISO"""
        
        for section in sections:
            # Extraire mots-clés ISO pertinents
            section_text_lower = (section.title + " " + section.content).lower()
            
            for category, keywords in self.iso_keywords.items():
                for keyword in keywords:
                    if keyword in section_text_lower:
                        section.keywords.append(category)
            
            # Déduplication
            section.keywords = list(set(section.keywords))
            
            # Déterminer le type de section
            title_lower = section.title.lower()
            if any(word in title_lower for word in ['formulaire', 'fiche', 'fo-', 'en-']):
                section.section_type = "formulaire"
            elif any(word in title_lower for word in ['procédure', 'processus', 'pr-']):
                section.section_type = "procedure"
            elif any(word in title_lower for word in ['instruction', 'mode opératoire', 'it-']):
                section.section_type = "instruction"
            else:
                section.section_type = "general"
        
        return sections
    
    def extract_and_segment(self, file_path: str) -> Tuple[List[DocumentSection], Dict]:
        """
        Pipeline complet : extraction + segmentation
        
        Returns:
            Tuple[List[DocumentSection], Dict]: (sections, métadonnées)
        """
        logger.info(f"Extraction du document: {file_path}")
        
        # Extraction
        text, metadata = self.extract_text(file_path)
        
        # Segmentation
        sections = self.segment_document(text, metadata)
        
        logger.info(f"Document segmenté en {len(sections)} sections")
        
        return sections, metadata


# Fonction utilitaire pour tester
def test_extractor(file_path: str):
    """Fonction de test"""
    extractor = DocumentExtractor()
    sections, metadata = extractor.extract_and_segment(file_path)
    
    print(f"\n=== Métadonnées ===")
    print(metadata)
    
    print(f"\n=== Sections détectées ({len(sections)}) ===")
    for i, section in enumerate(sections, 1):
        print(f"\n--- Section {i} ---")
        print(f"Titre: {section.title}")
        print(f"Numéro: {section.section_number}")
        print(f"Type: {section.section_type}")
        print(f"Page: {section.page_number}")
        print(f"Mots-clés ISO: {section.keywords}")
        print(f"Contenu (extrait): {section.content[:200]}...")


if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        test_extractor(sys.argv[1])
    else:
        print("Usage: python document_extractor.py <chemin_document>")